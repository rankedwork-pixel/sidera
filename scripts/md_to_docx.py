#!/usr/bin/env python3
"""Convert markdown files in docs/ to .docx format.

Uses python-docx to create properly formatted Word documents from
markdown source files. Handles headers, paragraphs, bullet lists,
tables, code blocks, and bold/italic text.

Usage:
    python3 scripts/md_to_docx.py
    python3 scripts/md_to_docx.py docs/skill-creation-guide.md
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor


def md_to_docx(md_path: Path, output_path: Path) -> None:
    """Convert a single markdown file to .docx."""
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Track state
    in_code_block = False
    code_lines: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block — flush
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                # Flush any pending table
                if in_table:
                    _add_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table detection (lines with |)
        if "|" in line and line.strip().startswith("|"):
            stripped = line.strip()
            # Skip separator rows (|---|---|)
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if cells:
                table_rows.append(cells)
                in_table = True
                i += 1
                continue

        # Flush pending table if we're past table lines
        if in_table:
            _add_table(doc, table_rows)
            table_rows = []
            in_table = False

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            # Add a thin line paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("_" * 60)
            run.font.color.rgb = RGBColor(180, 180, 180)
            run.font.size = Pt(8)
            i += 1
            continue

        # Headers
        if line.startswith("#"):
            level = len(line.split(" ")[0])  # count #'s
            header_text = line.lstrip("#").strip()
            if level == 1:
                doc.add_heading(header_text, level=0)
            elif level <= 4:
                doc.add_heading(header_text, level=min(level, 4))
            else:
                doc.add_heading(header_text, level=4)
            i += 1
            continue

        # Bullet lists (- or * or numbered)
        bullet_match = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)", line)
        if bullet_match:
            indent_level = len(bullet_match.group(1)) // 2
            content = bullet_match.group(3)
            is_numbered = bool(re.match(r"\d+\.", bullet_match.group(2)))

            if is_numbered:
                p = doc.add_paragraph(style="List Number")
            else:
                p = doc.add_paragraph(style="List Bullet")

            # Handle indent levels
            if indent_level > 0:
                p.paragraph_format.left_indent = Inches(0.25 * indent_level)

            _add_formatted_text(p, content)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        # Collect continuation lines (non-blank, non-special)
        para_lines = [line]
        j = i + 1
        while j < len(lines):
            next_line = lines[j]
            if (
                not next_line.strip()
                or next_line.startswith("#")
                or next_line.strip().startswith("```")
                or next_line.strip().startswith("- ")
                or next_line.strip().startswith("* ")
                or re.match(r"^\s*\d+\.\s+", next_line)
                or ("|" in next_line and next_line.strip().startswith("|"))
                or next_line.strip() in ("---", "***", "___")
            ):
                break
            para_lines.append(next_line)
            j += 1

        full_text = " ".join(ln.strip() for ln in para_lines)
        _add_formatted_text(p, full_text)
        i = j
        continue

    # Flush any remaining state
    if in_code_block and code_lines:
        _add_code_block(doc, code_lines)
    if in_table and table_rows:
        _add_table(doc, table_rows)

    doc.save(str(output_path))


def _add_formatted_text(paragraph, text: str) -> None:
    """Add text with bold/**text** and italic/*text* formatting."""
    # Split on bold (**text**) and italic (*text*) markers
    # Process bold first, then italic within each segment
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            # Bold text
            inner = part[2:-2]
            # Check for italic within bold
            italic_parts = re.split(r"(\*[^*]+\*)", inner)
            for ip in italic_parts:
                if ip.startswith("*") and ip.endswith("*"):
                    run = paragraph.add_run(ip[1:-1])
                    run.bold = True
                    run.italic = True
                else:
                    run = paragraph.add_run(ip)
                    run.bold = True
        else:
            # Check for italic
            italic_parts = re.split(r"(\*[^*]+\*)", part)
            for ip in italic_parts:
                if ip.startswith("*") and ip.endswith("*"):
                    run = paragraph.add_run(ip[1:-1])
                    run.italic = True
                else:
                    # Check for inline code `text`
                    code_parts = re.split(r"(`[^`]+`)", ip)
                    for cp in code_parts:
                        if cp.startswith("`") and cp.endswith("`"):
                            run = paragraph.add_run(cp[1:-1])
                            run.font.name = "Consolas"
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(80, 80, 80)
                        else:
                            if cp:
                                paragraph.add_run(cp)


def _add_code_block(doc: Document, lines: list[str]) -> None:
    """Add a code block as a formatted paragraph."""
    code_text = "\n".join(lines)
    if not code_text.strip():
        return

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(60, 60, 60)


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    """Add a table to the document."""
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Light Grid Accent 1"

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = ""
                p = cell.paragraphs[0]
                _add_formatted_text(p, cell_text)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

                # Bold the header row
                if i == 0:
                    for run in p.runs:
                        run.bold = True


def main() -> None:
    project_root = Path(__file__).parent.parent
    docs_dir = project_root / "docs"
    output_dir = docs_dir / "docx"
    output_dir.mkdir(exist_ok=True)

    # Determine which files to convert
    if len(sys.argv) > 1:
        # Convert specific files
        md_files = [Path(a) for a in sys.argv[1:]]
    else:
        # Convert all instruction docs (not progress log)
        md_files = [
            docs_dir / "skill-creation-guide.md",
            docs_dir / "architecture-decisions.md",
            docs_dir / "database-schema.md",
            docs_dir / "execution-flows.md",
            docs_dir / "adding-a-channel.md",
            docs_dir / "onboarding" / "01-Executive-Overview.md",
            docs_dir / "onboarding" / "02-How-It-Works.md",
            docs_dir / "onboarding" / "03-The-AI-Workforce.md",
            docs_dir / "onboarding" / "04-Cost-Estimates.md",
            docs_dir / "onboarding" / "05-Getting-Started.md",
        ]

    converted = 0
    for md_path in md_files:
        if not md_path.exists():
            print(f"  SKIP {md_path} (not found)")
            continue

        docx_name = md_path.stem + ".docx"
        # Preserve onboarding subfolder structure
        if "onboarding" in str(md_path):
            onboarding_dir = output_dir / "onboarding"
            onboarding_dir.mkdir(exist_ok=True)
            out_path = onboarding_dir / docx_name
        else:
            out_path = output_dir / docx_name

        try:
            md_to_docx(md_path, out_path)
            print(f"  OK   {md_path.name} -> {out_path.relative_to(project_root)}")
            converted += 1
        except Exception as exc:
            print(f"  FAIL {md_path.name}: {exc}")

    print(f"\nConverted {converted}/{len(md_files)} files to docs/docx/")


if __name__ == "__main__":
    main()
