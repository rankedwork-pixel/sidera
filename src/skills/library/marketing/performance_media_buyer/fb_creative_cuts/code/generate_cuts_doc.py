#!/usr/bin/env python3
"""
Generates a Word doc of Creative Cuts recommendations matching the exact
informal format from the Daily Stand Up notes.

No highlighting, no metrics unless needed to explain a nuance.
Short, informal, single-spaced between ad sets.

Usage:
    python3 generate_cuts_doc.py                  # uses latest CSV in data/
    python3 generate_cuts_doc.py path/to/file.csv  # uses specific CSV
"""

import os
import sys

from docx import Document
from docx.shared import Pt

sys.path.insert(0, os.path.dirname(__file__))
from collections import defaultdict

from creative_cuts import (
    MIN_AD_SET_SPEND,
    analyze_creatives_globally,
    dated_output_path,
    load_data,
    rank_ad_set,
    resolve_input_file,
)

# Max names to show per line (matching the informal standup style)
MAX_CUT_NAMES = 3
MAX_WATCH_NAMES = 2

# Only show cuts/watches that are actually worth calling out.
# If the waste score is below this, it's not worth mentioning — just N/C.
MIN_CUT_SPEND = 200  # Don't bother cutting ads with < $200 spend
MIN_WATCH_SPEND = 300  # Don't bother watching ads with < $300 spend

GENDER_ABBREV = {"female": "f", "male": "m"}
LANG_ABBREV = {"English": "eng", "Spanish": "span"}


def ad_set_label(state, gender, lang):
    g = GENDER_ABBREV.get(gender, gender[0])
    lang_abbr = LANG_ABBREV.get(lang, lang[:4].lower())
    return f"{state} {g}+{lang_abbr}"


def needs_context(ad):
    """
    Return a short parenthetical ONLY for genuinely conflicting signals.
    Two patterns:
      1. Bad CPL/internal but decent on-platform (ashleykelsey pattern)
      2. CPL-shielded: high CPBC but strong CPL — close call, worth flagging
    """
    cpbc = ad["on_platform_cpbc"]
    cpl = ad["cpl"]
    int_cpbc = ad["internal_cpbc"]

    # CPL-shielded ads: the analysis already flagged the conflict
    if ad.get("reason") == "CPL_SHIELDED" and ad.get("cpl_shield_note"):
        return f"${cpbc:,.0f} on platform, BUT ${cpl:,.0f} CPL — worth keeping?"

    # Conflicting signals: bad CPL/internal but decent on-platform
    if cpbc > 0 and cpl > 0 and int_cpbc > 0 and int_cpbc > cpbc * 3:
        return f"${cpl:,.0f} CPL, ${int_cpbc:,.0f} internal CPBC, BUT a ${cpbc:,.0f} on platform"

    return ""


def build_line(cuts, watches):
    """Build one informal ad set line, matching standup tone exactly."""
    # Filter to only meaningful recommendations
    real_cuts = [c for c in cuts if c["spend"] >= MIN_CUT_SPEND][:MAX_CUT_NAMES]
    real_watches = [w for w in watches if w["spend"] >= MIN_WATCH_SPEND][:MAX_WATCH_NAMES]

    if not real_cuts and not real_watches:
        return "N/C"

    cut_names = []
    for c in real_cuts:
        ctx = needs_context(c)
        if ctx:
            cut_names.append(f"{c['creative']} ({ctx})")
        else:
            cut_names.append(c["creative"])

    watch_names = [w["creative"] for w in real_watches]

    line = ""

    if cut_names:
        line += ", ".join(cut_names)

    if watch_names:
        if line:
            line += ", "
        line += "keep an eye on " + ", ".join(watch_names)

    # If only watches and no cuts
    if not cut_names and watch_names:
        line = "N/C, " + line

    return line


def main():
    # ─── Resolve input file ───────────────────────────────────────────────────
    cli_arg = sys.argv[1] if len(sys.argv) > 1 else None
    input_file = resolve_input_file(cli_arg)
    print(f"Input: {input_file}")

    records = load_data(input_file)
    global_flags = analyze_creatives_globally(records)

    ad_set_groups = defaultdict(list)
    for ad in records:
        ad_set_groups[(ad["state"], ad["gender"], ad["language"])].append(ad)

    results = {}
    for ad_set_key in sorted(ad_set_groups.keys()):
        ads = ad_set_groups[ad_set_key]
        total_spend = sum(a["spend"] for a in ads)
        if total_spend < MIN_AD_SET_SPEND:
            results[ad_set_key] = {"cuts": [], "watches": [], "skipped": True}
            continue
        cuts, watches, avg_cpbc, avg_cpl = rank_ad_set(ads, global_flags)
        results[ad_set_key] = {
            "cuts": cuts,
            "watches": watches,
            "avg_cpbc": avg_cpbc,
            "avg_cpl": avg_cpl,
            "skipped": False,
        }

    # ─── Build Word Doc ──────────────────────────────────────────────────────
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # Creative section header
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Creative:")
    run.bold = True
    run.font.size = Pt(11)

    # Launched from Preheat
    li = doc.add_paragraph(style="List Bullet")
    li.paragraph_format.space_after = Pt(0)
    li.paragraph_format.space_before = Pt(0)
    run = li.add_run("Launched from Preheat:")
    run.bold = True
    run.font.size = Pt(10)

    sub = doc.add_paragraph(style="List Bullet 2")
    sub.paragraph_format.space_after = Pt(0)
    sub.paragraph_format.space_before = Pt(0)
    sub.add_run("(none this run)").font.size = Pt(10)

    # Cuts header
    li = doc.add_paragraph(style="List Bullet")
    li.paragraph_format.space_after = Pt(0)
    li.paragraph_format.space_before = Pt(0)
    run = li.add_run("Cuts:")
    run.bold = True
    run.font.size = Pt(10)

    # Each ad set on one line, single spaced
    for ad_set_key in sorted(results.keys()):
        state, gender, lang = ad_set_key
        res = results[ad_set_key]
        label = ad_set_label(state, gender, lang)

        sub = doc.add_paragraph(style="List Bullet 2")
        sub.paragraph_format.space_after = Pt(0)
        sub.paragraph_format.space_before = Pt(0)

        run_label = sub.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.size = Pt(10)

        if res.get("skipped"):
            sub.add_run("N/C").font.size = Pt(10)
            continue

        line = build_line(res["cuts"], res["watches"])
        sub.add_run(line).font.size = Pt(10)

    # ─── Save with dated filename ─────────────────────────────────────────────
    output_doc = dated_output_path("Creative Cuts Recommendations", ".docx", records)
    doc.save(output_doc)
    print(f"Saved: {output_doc}")


if __name__ == "__main__":
    main()
